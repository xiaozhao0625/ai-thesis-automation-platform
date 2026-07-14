from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "AI论文自动化生产平台_UIUX_v1.2.2_工作流布局与任务导航校正说明.docx"


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    item = OxmlElement("w:shd")
    item.set(qn("w:fill"), fill)
    props.append(item)


def set_cell(cell, value, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(value)
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell(cell, value, bold=True, color=(255, 255, 255))
        shade(cell, "1B2639")
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            set_cell(cell, value)
    doc.add_paragraph()
    return table


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.1)
section.right_margin = Cm(2.1)
styles = doc.styles
styles["Normal"].font.name = "Microsoft YaHei"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Normal"].font.size = Pt(10.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("AI论文自动化生产平台\nUI/UX v1.2.2 工作流布局与任务导航校正说明")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(23, 37, 61)
subtitle = doc.add_paragraph("Spectrum Ledger｜研发交付校正包｜2026-07-15")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.color.rgb = RGBColor(86, 87, 217)
doc.add_paragraph()

heading(doc, "1. 本次校正结论")
doc.add_paragraph(
    "v1.2.2 在不改变既有视觉主题、双通道颜色、15 个页面范围、信息密度、ApprovalBar 层级和 v1.2.1 五条交互闭环的前提下，重建工作流展示与任务内导航。"
)
add_table(doc, ["范围", "交付结果", "验收口径"], [
    ["工作流画布", "数据驱动只读 DAG；HTML 节点层与 SVG 连线层分离。", "节点只出现一次，位于所属阶段内，边端点吸附端口。"],
    ["失效表达", "目录影响后，原位置切换 INVALIDATED；无底部重复节点面板。", "章节折叠组显示“5 个节点已失效”；质量闸门和交付包原位失效。"],
    ["任务导航", "九项 TaskSubNavigation 统一复用；面包屑可返回。", "切换、后退、前进、刷新、深链接均保持任务上下文。"],
    ["验收", "静态契约 + 7 条真实浏览器交互测试。", "既有四条闭环与新增画布/导航测试全部通过。"],
])

heading(doc, "2. 工作流画布实施规格")
doc.add_paragraph("画布采用独立世界坐标：阶段最小宽度 232px，节点宽度 180px，左右保留 26px 安全空间。V1.0 只读，支持横向滚动、空白区拖动、60%—160% 缩放和适配画布，但不支持用户修改节点位置。")
add_table(doc, ["对象", "实现契约"], [
    ["WorkflowNode", "绑定 node_id、display_name_cn、technical_id、stage_id、runtime_status、position、input_ports、output_ports。中文名称主显示，技术 ID 为次级信息。"],
    ["阶段", "启动、准备、规划、生产、工程、质量、交付；每列 232px，不允许节点跨列。"],
    ["边", "SVG path 从 source:out 到 target:in；普通、INVALIDATED、BLOCKED 使用实线/虚线/点线并保持不同语义。"],
    ["颜色", "领域顶边保持长期语义；例如 engineering_verify 始终工程绿 #16825D，运行状态仅通过紫色 Badge 表达。"],
])

heading(doc, "3. INVALIDATED 原位切换")
doc.add_paragraph("确认目录影响后，section_generate_ch3 至 section_generate_ch7、quality_gate_final 与 DeliveryPackage v3 在原 DAG 位置进入 INVALIDATED。领域顶边保留，叠加灰紫斜纹、中文失效状态和“目录 v4 影响”原因；受影响依赖边切换为失效线型。")
doc.add_paragraph("章节生成以折叠组承载：折叠时显示“章节生成 / 5 个节点已失效”；展开后显示五个章节节点。任一时刻，同一 node_id 在 DOM 中只存在一个实例。")

heading(doc, "4. TaskSubNavigation 与 URL 状态")
add_table(doc, ["项目", "规则"], [
    ["统一入口", "总览、工作流、项目与材料、资料与证据、目录、内容、工程验证、质量、交付。"],
    ["1280px", "二级导航横向滚动，质量和交付入口不隐藏。"],
    ["路由", "/tasks/task-001/{page}，并在 Query 中保留 taskId、node、attempt、claim、source、version、issue、package 等上下文。"],
    ["History", "点击使用 pushState/replaceState；popstate 还原页面与选中节点；刷新和复制深链接可恢复。"],
    ["面包屑", "“论文任务”返回任务列表；任务名返回总览；任意任务页可切回工作流。"],
])

heading(doc, "5. 自动化测试结果")
add_table(doc, ["编号", "场景", "结果"], [
    ["01", "目录确认后原位 INVALIDATED、无重复结果面板", "通过"],
    ["02", "DeadLetter 恢复创建独立 QUEUED NodeRun 并保留来源", "通过"],
    ["03", "质量复检从 RUNNING 至 quality_gate_final SUCCEEDED", "通过"],
    ["04", "终稿审批定位 v3，批准后启用正式下载", "通过"],
    ["05", "节点位于阶段内，所有 SVG 边端点落于端口容差内", "通过"],
    ["06", "章节组展开和折叠无重复 NodeRun", "通过"],
    ["07", "九项任务导航、后退/前进/刷新/深链接状态正确", "通过"],
])

heading(doc, "6. 交付物与限制")
doc.add_paragraph("代码仓库提交范围仅包括可复现 HTML 原型、静态契约测试、浏览器交互测试、校正说明和 handoff.md。原始资料库、交接素材、第三方 skill 和含敏感上下文的工作目录不上传。")
doc.add_paragraph("已进行浏览器级交互验证和 1440px 工作流画面检查。DOCX 结构与文本检查会执行；当前环境如无 LibreOffice，则 DOCX 的像素级渲染检查将被明确记录为未执行，而非伪报通过。")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("AI论文自动化生产平台 · UI/UX v1.2.2")
doc.save(OUT)
print(OUT)
